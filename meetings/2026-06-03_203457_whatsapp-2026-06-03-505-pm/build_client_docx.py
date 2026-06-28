#!/usr/bin/env python3
"""Build client-friendly Word docs from this meeting folder."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

FOLDER = Path(__file__).resolve().parent
GUIDE_PATH = FOLDER / "Antepartum Assessment - Study Guide.docx"
TRANSCRIPT_PATH = FOLDER / "Antepartum Assessment - Full Transcript.docx"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"


def add_bullet(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        run.bold = True
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)


def build_study_guide() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_heading("Antepartum Head-to-Toe Assessment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Study guide from your class recording")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Prepared: ").bold = True
    meta.add_run("June 3, 2026\n")
    meta.add_run("Recording length: ").bold = True
    meta.add_run("About 34 minutes\n")
    meta.add_run("Original file: ").bold = True
    meta.add_run("WhatsApp Audio 2026-06-03 at 5.05.29 PM")

    doc.add_heading("How to use this document", level=1)
    doc.add_paragraph(
        "This Word file is your main study guide. A second file in the same folder, "
        "“Antepartum Assessment - Full Transcript.docx,” has everything that was said, "
        "with timestamps. Open either file by double-clicking it."
    )
    doc.add_paragraph(
        "The transcript was created by computer from the audio. Some words may be "
        "wrong—always check with your instructor and your official course materials."
    )

    doc.add_heading("What this class was about", level=1)
    doc.add_paragraph(
        "This is a labor & delivery / antepartum nursing session: a full head-to-toe "
        "assessment on a pregnant patient. The instructor walks through each body system, "
        "what to look for, what to chart, and what questions to ask."
    )
    p = doc.add_paragraph()
    p.add_run("Most important idea: ").bold = True
    p.add_run(
        "Be thorough and accurate in your assessment and charting—including which side "
        "of the body you find something on. Vague notes can lead to wrong orders, missed "
        "problems, or sending someone home when it is not safe."
    )

    sections = [
        (
            "Neuro / general",
            [
                "Use your nursing assessment and judgment (for example, stroke)—not random internet answers.",
                "Remove all jewelry early, including hidden piercings. Metal is a problem if she needs emergency surgery.",
            ],
        ),
        (
            "Heart (cardiovascular)",
            ["Check pulses while you are assessing the upper body."],
        ),
        (
            "Lungs (respiratory)",
            [
                "Listen to the front and back of the chest. Sit the patient up—large breasts can hide lower lung sounds.",
                "Chart whether findings are on one side or both (for example, crackles on the right only).",
                "Ask about shortness of breath, smoking or vaping, asthma, and recent illness.",
                "In the hospital: no smoking (oxygen risk). Offer nicotine patch or gum. If she leaves to smoke, she may need to sign out AMA. Help with cravings so she stays on the unit.",
            ],
        ),
        (
            "Breasts",
            [
                "Still assess during pregnancy. Ask about breast implants or cosmetic surgery—patients often do not list these as surgery.",
                "Do a modified breast exam for lumps. Some women only see their OB, not a primary doctor—cancer can be found during pregnancy.",
                "Implants can affect breastfeeding depending on placement.",
            ],
        ),
        (
            "Abdomen / stomach & bowels",
            [
                "Bowel sounds move higher when she is pregnant—listen near the top of the uterus (fundus).",
                "Very active bowel sounds plus diarrhea: think hydration and risk of preterm labor.",
                "Belly pain: contractions are usually off-and-on; steady pain may be appendix, gallbladder, or other causes.",
                "Heartburn or chest discomfort can mimic a heart problem—ask and document clearly.",
            ],
        ),
        (
            "Bladder, kidneys & OB (GU)",
            [
                "Clean-catch urine: wipe, pee in the toilet, then catch the sample in the cup.",
                "UTI can harm the pregnancy, including preterm labor.",
                "For protein in urine (pre-eclampsia workup) when there is vaginal bleeding: use a straight catheter—not a clean catch, because blood can falsely raise protein.",
                "PROM = water broke at term. PPROM = membranes broke before labor when she is preterm.",
                "Placenta previa or bleeding history: nothing in the vagina (no sex; careful exams).",
                "Check vaginal discharge (color, smell, sores)—infection can cause preterm labor or sepsis.",
                "Water broken tests: know pooling and ferning. RomPlus, Amnisure, and litmus can be false positive with blood or KY jelly—swab before the internal exam when using litmus.",
            ],
        ),
        (
            "Back & spine",
            [
                "Back pain in pregnancy can be preterm labor—do not assume it is only normal back pain.",
                "Ask about scoliosis, back surgery, or injury—it affects epidural/spinal anesthesia. Tell anesthesia early when needed.",
            ],
        ),
        (
            "Legs, feet & skin",
            [
                "Document skin problems (rash, redness, sores) when you see them.",
                "Varicose veins, including in the vulva: higher DVT risk; rare heavy bleeding with pushing.",
                "Check foot pulses—feel left and right at the same time to compare.",
                "Painful, red calf: think blood clot (DVT).",
            ],
        ),
        (
            "Safety at home & feelings (psychosocial)",
            [
                "Even on a short visit, ask: Do you feel safe at home? Who will help you? Housing, money, partner issues?",
                "Example from class: homelessness found right before discharge—social work helped find a place to stay.",
            ],
        ),
    ]

    doc.add_heading("Study guide by body system", level=1)
    for heading, bullets in sections:
        doc.add_heading(heading, level=2)
        for b in bullets:
            add_bullet(doc, b)

    doc.add_heading("Tips for studying", level=1)
    for tip in [
        "Read this guide first, then use the Full Transcript file to find a topic (use Find: Ctrl+F).",
        "Search the transcript for words like: jewelry, crackles, AMA, straight cath, PROM, social worker.",
        "If a sentence in the transcript looks odd, trust your instructor—not the computer wording.",
    ]:
        add_bullet(doc, tip)

    doc.add_page_break()
    doc.add_heading("A note on accuracy", level=1)
    doc.add_paragraph(
        "This material is a study aid from an audio recording. It is not hospital policy. "
        "Use your program’s official books, skills checklists, and your clinical instructor "
        "for anything you do at the bedside."
    )

    doc.save(GUIDE_PATH)
    print("Wrote", GUIDE_PATH)


def build_transcript_doc() -> None:
    log = FOLDER / "transcript_log.txt"
    if not log.exists():
        return

    doc = Document()
    style_doc(doc)
    doc.add_heading("Antepartum Assessment — Full Transcript", 0)
    doc.add_paragraph(
        "Timestamped text from the recording. Times are MM:SS from the start of the audio."
    ).runs[0].italic = True
    doc.add_paragraph()

    for line in log.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line:
            continue
        if len(line) > 6 and line[2:3] == ":" and line[5:6] == " ":
            ts, _, rest = line.partition("  ")
            p = doc.add_paragraph()
            r = p.add_run(ts + "  ")
            r.bold = True
            p.add_run(rest)
        else:
            doc.add_paragraph(line)

    doc.save(TRANSCRIPT_PATH)
    print("Wrote", TRANSCRIPT_PATH)


if __name__ == "__main__":
    build_study_guide()
    build_transcript_doc()
