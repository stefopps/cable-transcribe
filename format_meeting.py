#!/usr/bin/env python3
"""Format a meeting folder into Word docs and optional client package."""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from cable_transcribe import (
    FINALIZE_TXT,
    LOG_FILE,
    MEETING_META_FILE,
    load_current_meeting,
)

SUMMARY_DOC = "Meeting Summary.docx"
TRANSCRIPT_DOC = "Full Transcript.docx"
PACKAGE_DIR = "Meeting-Package"


def _style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"


def _read_meeting_name(folder: Path) -> str:
    meta_path = folder / MEETING_META_FILE
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            name = str(meta.get("name", "")).strip()
            if name:
                return name
        except (json.JSONDecodeError, OSError):
            pass
    return folder.name.replace("_", " ")


def _parse_transcript_lines(raw: str) -> list[tuple[str, str]]:
    """Return (timestamp, text) pairs from transcript_log.txt."""
    rows: list[tuple[str, str]] = []
    for line in raw.splitlines():
        s = line.strip()
        if not s:
            continue
        m = re.match(r"^\[(\d{2}:\d{2})\]\s*(.*)$", s)
        if m:
            rows.append((m.group(1), m.group(2)))
            continue
        m2 = re.match(r"^(\d{2}:\d{2})\s{2,}(.*)$", s)
        if m2:
            rows.append((m2.group(1), m2.group(2)))
            continue
        if rows:
            ts, prev = rows[-1]
            rows[-1] = (ts, f"{prev} {s}".strip())
        else:
            rows.append(("", s))
    return rows


def build_transcript_doc(folder: Path, title: str) -> Path:
    log = folder / LOG_FILE
    if not log.exists():
        raise FileNotFoundError(f"No transcript: {log}")

    doc = Document()
    _style_doc(doc)
    doc.add_heading(f"{title} — Full Transcript", 0)
    doc.add_paragraph(
        "Timestamped text from the recording. Times are MM:SS from the start.",
    ).runs[0].italic = True
    doc.add_paragraph()

    for ts, text in _parse_transcript_lines(log.read_text(encoding="utf-8")):
        p = doc.add_paragraph()
        if ts:
            r = p.add_run(f"{ts}  ")
            r.bold = True
        p.add_run(text)

    out = folder / TRANSCRIPT_DOC
    doc.save(out)
    return out


def _add_finalize_line(doc: Document, line: str) -> None:
    s = line.strip()
    if not s:
        doc.add_paragraph()
        return
    u = s.upper()
    if u.startswith("GROUP:") or u.startswith("HIGHLIGHT:"):
        doc.add_heading(s, level=2)
        return
    if s.isupper() and len(s) > 3 and ":" not in s:
        doc.add_heading(s, level=1)
        return
    if s.startswith("-") or s.startswith("•"):
        doc.add_paragraph(s.lstrip("-• ").strip(), style="List Bullet")
        return
    if u.startswith("MOST_IMPORTANT:") or u.startswith("MOST IMPORTANT:"):
        p = doc.add_paragraph()
        p.add_run("Most important: ").bold = True
        p.add_run(s.split(":", 1)[-1].strip())
        return
    doc.add_paragraph(s)


def build_summary_doc(folder: Path, title: str) -> Path | None:
    fin = folder / FINALIZE_TXT
    if not fin.exists():
        return None

    doc = Document()
    _style_doc(doc)
    doc.add_heading(f"{title} — Meeting Summary", 0)
    sub = doc.add_paragraph("Generated from session transcript and notes")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    for line in fin.read_text(encoding="utf-8").splitlines():
        _add_finalize_line(doc, line)

    out = folder / SUMMARY_DOC
    doc.save(out)
    return out


def _find_recording(folder: Path) -> Path | None:
    meta_path = folder / MEETING_META_FILE
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            rec = str(meta.get("recording_file", "")).strip()
            if rec and (folder / rec).is_file():
                return folder / rec
        except (json.JSONDecodeError, OSError):
            pass
    for pattern in ("recording.*", "*.mp4", "*.mp3", "*.m4a", "*.wav"):
        hits = list(folder.glob(pattern))
        if hits:
            return hits[0]
    return None


def build_client_package(folder: Path, title: str) -> Path:
    """Self-contained folder with recording, Word docs, and double-click shortcuts."""
    pkg = folder / PACKAGE_DIR
    if pkg.exists():
        shutil.rmtree(pkg)
    pkg.mkdir(parents=True)

    summary = folder / SUMMARY_DOC
    transcript = folder / TRANSCRIPT_DOC
    recording = _find_recording(folder)

    if summary.exists():
        shutil.copy2(summary, pkg / summary.name)
    if transcript.exists():
        shutil.copy2(transcript, pkg / transcript.name)
    if recording:
        shutil.copy2(recording, pkg / f"Recording{recording.suffix.lower()}")

    start = pkg / "START HERE.txt"
    start.write_text(
        f"{title.upper()}\n"
        f"{'=' * len(title)}\n\n"
        "Everything is in this folder. Copy the whole folder to another PC\n"
        "(or unzip the ZIP Steve sent). No install needed to read the Word files.\n\n"
        "FILES\n"
        "-----\n"
        + (f"  {summary.name}\n      Meeting summary — start here.\n\n" if summary.exists() else "")
        + (f"  {transcript.name}\n      Full timestamped transcript (Ctrl+F to search).\n\n" if transcript.exists() else "")
        + (f"  Recording{recording.suffix.lower() if recording else ''}\n      Original audio/video.\n\n" if recording else "")
        + "SHORTCUTS\n"
        "---------\n"
        "  Double-click the .bat files to open docs or play the recording.\n",
        encoding="utf-8",
    )

    if summary.exists():
        (pkg / "1 - OPEN SUMMARY.bat").write_text(
            f'@echo off\nstart "" "{summary.name}"\n',
            encoding="utf-8",
        )
    if transcript.exists():
        (pkg / "2 - OPEN TRANSCRIPT.bat").write_text(
            f'@echo off\nstart "" "{transcript.name}"\n',
            encoding="utf-8",
        )
    if recording:
        rec_name = f"Recording{recording.suffix.lower()}"
        (pkg / "3 - PLAY RECORDING.bat").write_text(
            f'@echo off\nstart "" "{rec_name}"\n',
            encoding="utf-8",
        )

    return pkg


def format_meeting_folder(
    folder: Path,
    *,
    package: bool = False,
) -> dict[str, Path]:
    """Build Word docs (and optional client package) for one meeting folder."""
    folder = folder.resolve()
    if not folder.is_dir():
        raise FileNotFoundError(f"Not a folder: {folder}")

    title = _read_meeting_name(folder)
    results: dict[str, Path] = {}

    print(f"Formatting: {title}")
    results["transcript"] = build_transcript_doc(folder, title)
    print(f"  Wrote {results['transcript'].name}")

    summary = build_summary_doc(folder, title)
    if summary:
        results["summary"] = summary
        print(f"  Wrote {summary.name}")
    else:
        print(f"  (no {FINALIZE_TXT} — summary doc skipped)")

    if package:
        pkg = build_client_package(folder, title)
        results["package"] = pkg
        print(f"  Package: {pkg}")

    return results


def resolve_folder(path: str | None) -> Path:
    if path:
        return Path(path)
    loaded = load_current_meeting()
    if loaded is not None:
        return loaded.folder
    raise SystemExit("No meeting folder given and no current_meeting.json found.")


def main() -> None:
    parser = argparse.ArgumentParser(description="Format meeting folder to Word docs")
    parser.add_argument(
        "folder",
        nargs="?",
        help="Meeting folder (default: current meeting from current_meeting.json)",
    )
    parser.add_argument(
        "--package",
        action="store_true",
        help="Also build Meeting-Package/ with bats + START HERE.txt",
    )
    args = parser.parse_args()

    try:
        folder = resolve_folder(args.folder)
        format_meeting_folder(folder, package=args.package)
    except FileNotFoundError as exc:
        print(exc, file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
